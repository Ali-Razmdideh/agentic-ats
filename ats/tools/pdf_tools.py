from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import urlparse

from claude_agent_sdk import tool

log = logging.getLogger("ats.tools.pdf_tools")


def _normalize_url(href: str) -> str | None:
    """Return ``href`` if it looks like a real http(s) / mailto URL, else None."""
    if not href:
        return None
    href = href.strip()
    if not href:
        return None
    if href.startswith("mailto:") or href.startswith("tel:"):
        return href
    parsed = urlparse(href)
    if parsed.scheme in ("http", "https") and parsed.netloc:
        return href
    return None


def _extract_pdf_links(path: Path) -> list[str]:
    """Pull URI targets from /Link annotations on every PDF page.

    PDFs commonly render hyperlinks as styled visible text with the real
    target stored as a separate annotation. ``page.extract_text()`` only
    returns the visible text, dropping the URL. Walking the annotations
    fills that gap so the parser can populate ``links``.
    """
    try:
        from pypdf import PdfReader
        from pypdf.generic import IndirectObject
    except Exception:  # pragma: no cover - pypdf always installed in our env
        return []

    out: list[str] = []
    seen: set[str] = set()
    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # malformed file
        log.debug("pdf reader failed", extra={"err": str(exc)})
        return []

    for page in reader.pages:
        annots = page.get("/Annots") or []
        if isinstance(annots, IndirectObject):
            try:
                annots = annots.get_object() or []
            except Exception:
                annots = []
        for ann in annots:
            try:
                obj = ann.get_object() if isinstance(ann, IndirectObject) else ann
            except Exception:
                continue
            if not hasattr(obj, "get"):
                continue
            if obj.get("/Subtype") != "/Link":
                continue
            action = obj.get("/A")
            if isinstance(action, IndirectObject):
                try:
                    action = action.get_object()
                except Exception:
                    continue
            if not action or not hasattr(action, "get"):
                continue
            uri = action.get("/URI")
            if not isinstance(uri, str):
                # pypdf sometimes wraps strings in TextStringObject; coerce.
                uri = str(uri) if uri is not None else ""
            normalized = _normalize_url(uri)
            if normalized and normalized not in seen:
                seen.add(normalized)
                out.append(normalized)
    return out


def _extract_docx_links(path: Path) -> list[str]:
    """Pull hyperlink targets from a .docx by walking the document relationships."""
    try:
        import docx  # python-docx
    except Exception:  # pragma: no cover
        return []

    out: list[str] = []
    seen: set[str] = set()
    try:
        d = docx.Document(str(path))
    except Exception:
        return []
    rels = getattr(d.part, "rels", {}) or {}
    for rel in rels.values():
        if getattr(rel, "reltype", "").endswith("/hyperlink"):
            target = getattr(rel, "target_ref", "") or ""
            normalized = _normalize_url(target)
            if normalized and normalized not in seen:
                seen.add(normalized)
                out.append(normalized)
    return out


def _format_links_section(urls: Iterable[str]) -> str:
    """Format a LINKS block to append to the extracted text."""
    items = list(urls)
    if not items:
        return ""
    body = "\n".join(items)
    return f"\n\nLINKS:\n{body}\n"


def extract_text_from_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        body = "\n".join((page.extract_text() or "") for page in reader.pages)
        links = _extract_pdf_links(path)
        return body + _format_links_section(links)
    if suffix == ".docx":
        import docx  # python-docx

        d = docx.Document(str(path))
        body = "\n".join(p.text for p in d.paragraphs)
        links = _extract_docx_links(path)
        return body + _format_links_section(links)
    raise ValueError(f"Unsupported file type: {suffix}")


@tool(
    "read_resume",
    "Extract plain text from a resume file (.pdf, .docx, .txt, .md). "
    "Hyperlink URLs from the document are appended in a LINKS: section at "
    "the end so the parser can populate the candidate's links list.",
    {"path": str},
)
async def read_resume(args: dict[str, Any]) -> dict[str, Any]:
    path = Path(args["path"])
    if not path.exists():
        return {
            "content": [{"type": "text", "text": f"ERROR: file not found: {path}"}],
            "isError": True,
        }
    try:
        text = extract_text_from_path(path)
    except Exception as exc:
        return {
            "content": [{"type": "text", "text": f"ERROR: {exc}"}],
            "isError": True,
        }
    return {"content": [{"type": "text", "text": text}]}
